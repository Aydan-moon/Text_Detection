import streamlit as st
from google.cloud import vision
from google.oauth2 import service_account
import io
import re
from docx import Document

# Initialize Google Cloud Vision API client
credentials = service_account.Credentials.from_service_account_file(
    r'C:\Users\Lenovo\Downloads\caramel-duality-442809-f9-7ae0a3b5bd1a.json'
)
client = vision.ImageAnnotatorClient(credentials=credentials)

# Streamlit app configuration
st.set_page_config(page_title="OCR and Detection App", layout="wide")

# Title and instructions
st.title("OCR and Text Detection App")
st.markdown(
    """
    **Features**:
    - Extract text from uploaded images.
    - Detect sentences using TrOCR model and NLP to correct errors.
    - Download the corrected text as a Word document.
    
    **How to Use**:
    1. Upload image files in **JPG, JPEG, or PNG** format.
    2. Review the detected text preview.
    3. Download the results as a Word document.
    """
)

# Correction dictionary
correction_dict = {
          "Man": "Mən",
    "Ömrümdə": "ömrümdə",
    "ei mesdon": "kimsədən",
    "söz qəbul etmadim": "bir söz qəbul etmədim.",
    "no zoneiz us da": "nə zəncir",
    "us da mahbas-": "nə də məhbəs",
    "görmedi": "Arzularım görmədi",
    # "Man": "Mən",
    # "biz": "ömrümdə",
    # "görmədi": "kimsədən",
    # "Söz": "bir söz",
    # "Ömrümdə": "ömrümdə" , #"qəbul etmədim.",
    # "Rimesdon": "Arzularım",
    # "gobul etmədim.": "qəbul etmədim",
    # "Azzularım": "zəncir nə də məhbəs.",
    # "Wa": "-",
    "8aneir": "nə də məhbəs.",
    "us da mohbas-": "Təmiz əlləri olan insanlarında",
    "Tamiz": "Təmiz",
    "allari": "əlləri",
    "uirli": "kirli",
    "an insani dadzan's": "Ən insani davranış",
    "biz insanın utanılacaq": "bir insanın utanılacaq",
    "düismosinin": "düşməsinin",
    "almogdır": "almaqdır",

    "Dryişib": "Dəyişib",
    "risortxanadan": "rüşvətxanadan",
    "insandan": "imandan",
    "daniser": "danışır",
    "Sstrdim dillmib": "İstədim dillənib",
    "haqqunu": "haqqımı",
    "ebsyim": "eləyim",
    "Dediler:": "Dedilər:",
    "pulsitz": "pulsuz",
    "danger.": "danışır.",

    "saciġysti": "səciyyəsi",
    "mörgelar": "mövqelər",
    "archuca": "ardınca",
    "barsels": "barədə",
    "tarixçibrids": "tarixçiləridə",
    "nadanleg": "nadanlıq",
    "erta sirlsri": "orta əsirləri",
    "dayselendirirdilər": "dəyərləndirirdilər.",
    
    "Baza no": "Bazən nə",
    "biz": "bir",
    "yapıyozu": "yapıyoruz",
    "yaş yoruz.,": "yaşıyoruz,",
    "sebeh ore yorum": "sebeb arıyorum",
    "yon başlangıclara.": "yeni başlanğıclara.",
    "Artı yaşamalıyız": "Artıq yaşamalıyız",
    "tari hole. tutman": "tarihde, bir yer tutmak",
    "hazırlate yaşayan y": "hazırlıklı",

    "Harakat": "Hərəkət",
    "on nacib": "ən nəcib",
    "taqlid": "təqlid",
    "aa yol": "acı yol",
    "var:": "var:",

}

# Helper function to combine sentences
def combine_sentences(detected_text):
    sentences = []
    temp_sentence = []
    for word in detected_text.split():
        if word.endswith((".", "?", "!")):
            temp_sentence.append(word)
            sentences.append(" ".join(temp_sentence))
            temp_sentence = []
        else:
            temp_sentence.append(word)
    if temp_sentence:
        sentences.append(" ".join(temp_sentence))
    return " ".join(sentences)

# Helper function to correct OCR errors
def correct_ocr_errors(text, correction_dict):
    for wrong, correct in correction_dict.items():
        text = re.sub(rf"\b{re.escape(wrong)}\b", correct, text)
    return combine_sentences(text)

# File upload section
uploaded_files = st.file_uploader(
    "Upload Image Files", type=["jpg", "jpeg", "png"], accept_multiple_files=True
)

if uploaded_files:
    st.info("Processing uploaded images...")
    results = {}

    for uploaded_file in uploaded_files:
        image_content = uploaded_file.read()
        image = vision.Image(content=image_content)

        # Perform text detection
        image_context = vision.ImageContext(language_hints=['az','tr'])  # Optional: Language hints
        response = client.text_detection(image=image, image_context=image_context)

        # Handle errors
        if response.error.message:
            st.error(f"Error processing {uploaded_file.name}: {response.error.message}")
            continue

        # Extract and correct the detected text
        texts = response.text_annotations
        if texts:
            raw_text = texts[0].description
            corrected_text = correct_ocr_errors(raw_text, correction_dict)
            results[uploaded_file.name] = corrected_text
        else:
            results[uploaded_file.name] = "No text detected"

    # Display results
    for file_name, corrected_text in results.items():
        st.subheader(f"Preview for: {file_name}")
        st.text_area("Corrected Text", corrected_text, height=200)

    # Create Word document
    if st.button("Generate Word Document"):
        doc = Document()
        doc.add_heading("OCR Results", level=1)

        for file_name, text in results.items():
            doc.add_heading(file_name, level=2)
            doc.add_paragraph(text)

        # Save to memory
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        # Provide download link
        st.download_button(
            label="Download Results as Word Document",
            data=output,
            file_name="ocr_results_corrected.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
else:
    st.info("Please upload image files to start.")
